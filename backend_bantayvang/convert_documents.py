#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script để chuyển đổi file .docx thành text và đọc file .txt
Yêu cầu: pip install python-docx
"""

import os
import sys
from pathlib import Path

def install_requirements():
    """Cài đặt các thư viện cần thiết"""
    try:
        import docx
        print("✓ python-docx đã được cài đặt")
    except ImportError:
        print("Đang cài đặt python-docx...")
        os.system("pip install python-docx")
        try:
            import docx
            print("✓ Cài đặt python-docx thành công")
        except ImportError:
            print("❌ Không thể cài đặt python-docx. Vui lòng chạy: pip install python-docx")
            return False
    return True

def read_docx_file(file_path):
    """Đọc nội dung file .docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        content = []
        
        print(f"\n📄 Đang đọc file: {file_path}")
        print("=" * 60)
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Đọc tables nếu có
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
        
    except Exception as e:
        print(f"❌ Lỗi khi đọc file .docx: {e}")
        return None

def read_txt_file(file_path):
    """Đọc nội dung file .txt"""
    try:
        print(f"\n📄 Đang đọc file: {file_path}")
        print("=" * 60)
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        return content
        
    except Exception as e:
        print(f"❌ Lỗi khi đọc file .txt: {e}")
        return None

def save_converted_content(content, output_file):
    """Lưu nội dung đã chuyển đổi"""
    try:
        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(content)
        print(f"✓ Đã lưu nội dung vào: {output_file}")
    except Exception as e:
        print(f"❌ Lỗi khi lưu file: {e}")

def main():
    """Hàm chính"""
    print("🔄 Bắt đầu chuyển đổi tài liệu...")
    
    # Kiểm tra và cài đặt thư viện
    if not install_requirements():
        return
    
    # Đường dẫn thư mục
    folder_path = Path("decuongvayeucau")
    
    if not folder_path.exists():
        print(f"❌ Không tìm thấy thư mục: {folder_path}")
        return
    
    # Tìm và xử lý các file
    docx_files = list(folder_path.glob("*.docx"))
    txt_files = list(folder_path.glob("*.txt"))
    
    print(f"\n📁 Tìm thấy:")
    print(f"   - {len(docx_files)} file .docx")
    print(f"   - {len(txt_files)} file .txt")
    
    # Xử lý file .docx
    for docx_file in docx_files:
        print(f"\n🔄 Xử lý file: {docx_file.name}")
        content = read_docx_file(docx_file)
        
        if content:
            print(content[:500] + "..." if len(content) > 500 else content)
            
            # Lưu vào file text
            output_file = folder_path / f"{docx_file.stem}_converted.txt"
            save_converted_content(content, output_file)
    
    # Xử lý file .txt
    for txt_file in txt_files:
        print(f"\n🔄 Xử lý file: {txt_file.name}")
        content = read_txt_file(txt_file)
        
        if content:
            print(content)
    
    print("\n✅ Hoàn thành chuyển đổi tài liệu!")

if __name__ == "__main__":
    main()